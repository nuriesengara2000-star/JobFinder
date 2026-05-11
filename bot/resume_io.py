"""Extract resume text from uploaded files (PDF / DOCX / TXT)."""

from __future__ import annotations

import io
import logging

LOG = logging.getLogger(__name__)

MAX_BYTES = 5 * 1024 * 1024  # 5 MB cap on uploaded resumes


class ResumeReadError(Exception):
    """Raised when the uploaded file cannot be parsed into text."""


def extract_text(data: bytes, filename: str | None = None, mime_type: str | None = None) -> str:
    """Return plain text extracted from a resume file.

    Format is detected from filename extension first, then mime_type. Raises
    ``ResumeReadError`` with a user-friendly message on any failure.
    """
    if not data:
        raise ResumeReadError("Файл пустой.")
    if len(data) > MAX_BYTES:
        raise ResumeReadError(f"Файл слишком большой (>{MAX_BYTES // (1024 * 1024)} MB).")

    fmt = _detect_format(filename, mime_type, data)
    LOG.info("Extracting resume: format=%s filename=%r size=%d", fmt, filename, len(data))

    if fmt == "pdf":
        return _from_pdf(data)
    if fmt == "docx":
        return _from_docx(data)
    if fmt == "txt":
        return _from_txt(data)

    raise ResumeReadError(
        "Поддерживаются только PDF, DOCX и TXT. "
        "Если у тебя .doc — сохрани его как .docx."
    )


def _detect_format(filename: str | None, mime: str | None, data: bytes) -> str | None:
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return "pdf"
    if name.endswith(".docx"):
        return "docx"
    if name.endswith(".txt"):
        return "txt"

    m = (mime or "").lower()
    if "pdf" in m:
        return "pdf"
    if "wordprocessingml" in m or "officedocument.wordprocessing" in m:
        return "docx"
    if m.startswith("text/"):
        return "txt"

    # Magic-byte fallbacks
    if data[:4] == b"%PDF":
        return "pdf"
    if data[:4] == b"PK\x03\x04":  # any zip — DOCX is a zip
        return "docx"
    return None


def _from_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ResumeReadError("Сервер: не установлен pypdf.") from exc

    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001 — pypdf raises a variety of errors
        raise ResumeReadError(f"Не получилось открыть PDF: {exc}") from exc

    if reader.is_encrypted:
        try:
            reader.decrypt("")  # try empty-password decrypt
        except Exception as exc:
            raise ResumeReadError("PDF зашифрован паролем — расшифруй и пришли снова.") from exc

    pages: list[str] = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception as exc:  # noqa: BLE001
            LOG.warning("PDF page extract failed: %s", exc)
            continue

    text = "\n".join(p.strip() for p in pages if p.strip())
    if not text.strip():
        raise ResumeReadError(
            "В PDF не нашлось текстового слоя — похоже это скан. "
            "Пришли DOCX или текстовую версию резюме."
        )
    return text


def _from_docx(data: bytes) -> str:
    try:
        from docx import Document  # python-docx
    except ImportError as exc:
        raise ResumeReadError("Сервер: не установлен python-docx.") from exc

    try:
        doc = Document(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        raise ResumeReadError(f"Не получилось открыть DOCX: {exc}") from exc

    parts: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = (cell.text or "").strip()
                if t:
                    parts.append(t)

    text = "\n".join(parts)
    if not text.strip():
        raise ResumeReadError("DOCX оказался пустым.")
    return text


def _from_txt(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "cp1251", "cp1252", "latin-1"):
        try:
            text = data.decode(encoding)
            return text
        except UnicodeDecodeError:
            continue
    raise ResumeReadError("Не удалось определить кодировку TXT-файла.")
